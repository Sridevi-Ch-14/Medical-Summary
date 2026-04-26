"""
ClinIQ - Word Document Extractor (Phase 1)
============================================
Extracts text from .docx files including paragraphs and tables.
Uses python-docx for structured extraction.
"""

import re
import logging
from pathlib import Path

logger = logging.getLogger(__name__)


def extract_text_from_docx(file_path: str) -> str:
    """
    Extract all text from a Word (.docx) document.
    
    Handles:
    - Regular paragraphs
    - Tables (flattened to row-based text with column alignment)
    - Headers and footers (optional)
    
    Args:
        file_path: Path to the .docx file
        
    Returns:
        Cleaned extracted text
    """
    file_path = str(Path(file_path).resolve())
    logger.info(f"Extracting text from DOCX: {file_path}")

    try:
        from docx import Document
    except ImportError:
        logger.error("python-docx not installed. Run: pip install python-docx")
        return ""

    try:
        doc = Document(file_path)
        text_parts = []

        # Extract all paragraphs
        for para in doc.paragraphs:
            stripped = para.text.strip()
            if stripped:
                text_parts.append(stripped)

        # Extract all tables
        for table_idx, table in enumerate(doc.tables):
            table_lines = []
            for row in table.rows:
                cells = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    cells.append(cell_text if cell_text else "")
                row_text = "    ".join(cells)
                if row_text.strip():
                    table_lines.append(row_text)
            
            if table_lines:
                text_parts.append("\n".join(table_lines))

        raw_text = "\n".join(text_parts)
        cleaned = _clean_docx_text(raw_text)
        logger.info(f"Extracted {len(cleaned)} characters from DOCX")
        return cleaned

    except Exception as e:
        logger.error(f"DOCX extraction failed: {e}")
        return ""


def _clean_docx_text(raw_text: str) -> str:
    """Clean extracted DOCX text — remove noise and normalize whitespace."""
    text = raw_text

    # Remove non-printable characters (keep standard ASCII + newlines)
    text = re.sub(r'[^\x20-\x7E\n\t]', '', text)

    # Collapse multiple blank lines
    text = re.sub(r'\n{3,}', '\n\n', text)

    # Collapse multiple spaces
    text = re.sub(r' {2,}', '  ', text)

    return text.strip()
